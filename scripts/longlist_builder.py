"""Excel + docx builder for Herb's Phase 3 (longlist drafts) and Phase 5 (final report).

Outputs are written to runs/[slug]/ in the repo so the team can browse them directly.

Public API:
    build_longlist_v1(slug, rows)             -> Path to longlist-v1.xlsx
    build_longlist_vN(slug, n, new_rows)      -> Path to longlist-v{n}.xlsx (preserves prior order)
    parse_scorecard(text)                     -> dict (keys: score, recommendation, gates, lps, ...)
    build_final_longlist(slug, scorecards)    -> Path to final-longlist-[slug].xlsx
    build_final_summary_docx(slug, summary)   -> Path to final-summary-[slug].docx
    tag_comp_focus(tech, sectors)             -> "HIGH" | "MED" | "LOW"

Each row dict for build_longlist_v1 / build_longlist_vN should have the LONGLIST_COLUMNS keys
below; missing keys default to "" (Unknown).
"""
from __future__ import annotations
import re
from pathlib import Path
from typing import Iterable

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

REPO_ROOT = Path(__file__).resolve().parent.parent
RUNS_DIR = REPO_ROOT / "runs"

# ---- column schema (matches today's longlist-v1.xlsx output verbatim) ----
LONGLIST_COLUMNS = [
    "Company", "Domain", "HQ Country", "Stage", "Raised ($M)", "Last Round", "Investors",
    "Tech (1 line)", "Sectors served", "Source", "Pipedrive Status", "Pre-screen",
]
COMP_FOCUS_COL = "Comp Focus"   # added on v2 and later

SCORE_HEADERS = [
    "Company", "Domain", "Score /5", "Recommendation",
    "Gate: Stage/Rev", "Gate: Sector", "Gate: Tech", "Gate: Climate",
    "Gate: Geography", "Gate: Biz model", "Gate: LP fit",
    "LP: Nouryon", "LP: Bühler", "LP: FrieslandCampina",
    "Top Question", "Analyst Note",
]

# ---- styling ----
GREEN = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
AMBER = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
RED   = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
GREY  = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
BLUE  = PatternFill(start_color="DDEBF7", end_color="DDEBF7", fill_type="solid")
BOLD  = Font(bold=True)
WRAP  = Alignment(wrap_text=True, vertical="top")

# ============================================================================
# Comp Focus tagging (lifted from today's _build_v2.py)
# ============================================================================
_HIGH_KEYWORDS = [
    "ai protein", "ai enzyme", "ai-driven enzyme", "ai-driven protein", "ai for enzyme", "ai for protein",
    "ml ", "machine learning", "deep learning", "generative ai", "generative model",
    "computational", "in silico", "in-silico", "language model", "esm ", "alphafold", "rfdiffusion",
    "proteinmpnn", "diffusion model", "molecular dynamics", "quantum", "data platform",
    "bioinformatics for", "molecular modeling", "design platform", "genomics", "ai analytics",
    "ai green", "ai retrosynthesis", "ai robotics", "ai biology", "ai-guided", "ai-powered",
    "foundation model", "foundation ai", "protein language", "protein design platform",
    "engineering software", "process optimization software", "trrosetta", "physics-ai",
    "physics-based", "point-cloud ai", "pro-prime",
]
_MED_KEYWORDS = [
    "directed evolution", "high throughput screening", "ultra-high throughput", "microfluidic",
    "synthetic biology", "metabolic engineering", "biocatalyst engineering", "amino acid",
    "novel enzyme design", "protein engineering for", "enzyme discovery",
]


def tag_comp_focus(tech: str | None, sectors: str | None = "") -> str:
    """Return HIGH | MED | LOW based on tech/sectors keywords."""
    s = ((tech or "") + " " + (sectors or "")).lower()
    if any(k in s for k in _HIGH_KEYWORDS):
        return "HIGH"
    if any(k in s for k in _MED_KEYWORDS):
        return "MED"
    return "LOW"


def _fill_for_focus(tag: str) -> PatternFill | None:
    return {"HIGH": GREEN, "MED": AMBER, "LOW": RED}.get(tag)


# ============================================================================
# Sheet construction helpers
# ============================================================================
def _row_to_values(row: dict, columns: list[str]) -> list:
    return [row.get(c, "") for c in columns]


def _write_header(ws, headers: list[str]) -> None:
    for j, h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=j, value=h)
        c.font = BOLD
        c.alignment = WRAP


def _autosize_basic(ws, n_cols: int) -> None:
    for j in range(1, n_cols + 1):
        ws.column_dimensions[get_column_letter(j)].width = 22


def _placeholder_score_sheet(wb: Workbook) -> None:
    """Sheet 2 placeholder used in v1/vN until Phase 5 populates it."""
    if "Icos Fit Scores" in wb.sheetnames:
        del wb["Icos Fit Scores"]
    ws = wb.create_sheet("Icos Fit Scores")
    _write_header(ws, SCORE_HEADERS)
    ws.cell(row=2, column=1, value=(
        "Reply to Herb's email with 'Score rows X, Y, Z' (or company names) to trigger "
        "Icos Fit evaluations on selected companies. Sheet 2 populates in the final report."
    )).alignment = WRAP
    for j in range(1, len(SCORE_HEADERS) + 1):
        ws.column_dimensions[get_column_letter(j)].width = 18


def _ensure_run_dir(slug: str) -> Path:
    d = RUNS_DIR / slug
    d.mkdir(parents=True, exist_ok=True)
    return d


# ============================================================================
# Build longlist v1 — Phase 3 first round
# ============================================================================
def build_longlist_v1(slug: str, rows: Iterable[dict]) -> Path:
    """Create runs/[slug]/longlist-v1.xlsx from the post-dedup, post-Pipedrive-cross-check rows."""
    d = _ensure_run_dir(slug)
    out = d / "longlist-v1.xlsx"

    rows = list(rows)
    wb = Workbook()
    ws = wb.active
    ws.title = "Longlist"

    _write_header(ws, LONGLIST_COLUMNS)
    for i, row in enumerate(rows, start=2):
        for j, val in enumerate(_row_to_values(row, LONGLIST_COLUMNS), start=1):
            ws.cell(row=i, column=j, value=val).alignment = WRAP

    _autosize_basic(ws, len(LONGLIST_COLUMNS))
    _placeholder_score_sheet(wb)
    wb.save(out)
    return out


# ============================================================================
# Build longlist v2/v3 — Phase 4 iteration
# ============================================================================
def build_longlist_vN(slug: str, n: int, new_rows: Iterable[dict]) -> Path:
    """Read runs/[slug]/longlist-v{n-1}.xlsx, append new_rows, ensure Comp Focus column.

    Preserves existing row order so author references ("score row 5") stay valid.
    new_rows are appended at the bottom (tagged HIGH by default — they're round-N additions).
    """
    if n < 2:
        raise ValueError(f"build_longlist_vN requires n>=2; got {n}")
    d = _ensure_run_dir(slug)
    src = d / f"longlist-v{n-1}.xlsx"
    out = d / f"longlist-v{n}.xlsx"
    if not src.is_file():
        raise FileNotFoundError(f"prior longlist not found: {src}")

    wb = load_workbook(src)
    ws = wb["Longlist"]

    # Determine current header & ensure Comp Focus column exists
    headers = [c.value for c in ws[1]]
    if COMP_FOCUS_COL not in headers:
        comp_idx = len(headers) + 1
        ws.cell(row=1, column=comp_idx, value=COMP_FOCUS_COL).font = BOLD
        # Tag existing rows
        tech_idx = headers.index("Tech (1 line)") + 1
        sect_idx = headers.index("Sectors served") + 1
        pre_idx  = headers.index("Pre-screen") + 1
        for r in range(2, ws.max_row + 1):
            tech = ws.cell(row=r, column=tech_idx).value or ""
            sectors = ws.cell(row=r, column=sect_idx).value or ""
            pre = (ws.cell(row=r, column=pre_idx).value or "").lower()
            tag = "—" if pre.startswith("fail") else tag_comp_focus(tech, sectors)
            cell = ws.cell(row=r, column=comp_idx, value=tag)
            fill = _fill_for_focus(tag)
            if fill:
                cell.fill = fill
        ws.column_dimensions[get_column_letter(comp_idx)].width = 12
    else:
        comp_idx = headers.index(COMP_FOCUS_COL) + 1

    # Append new rows
    start_new = ws.max_row + 1
    new_rows = list(new_rows)
    for i, row in enumerate(new_rows):
        r = start_new + i
        for j, col in enumerate(LONGLIST_COLUMNS, start=1):
            ws.cell(row=r, column=j, value=row.get(col, "")).alignment = WRAP
        # Tag and color
        tag = row.get(COMP_FOCUS_COL) or tag_comp_focus(row.get("Tech (1 line)", ""),
                                                       row.get("Sectors served", ""))
        cell = ws.cell(row=r, column=comp_idx, value=tag)
        fill = _fill_for_focus(tag)
        if fill:
            cell.fill = fill
        # Light-blue tint on Source col so authors see what's new at a glance
        ws.cell(row=r, column=LONGLIST_COLUMNS.index("Source") + 1).fill = BLUE

    # Sheet 2 stays placeholder unless build_final_longlist runs
    if "Icos Fit Scores" not in wb.sheetnames:
        _placeholder_score_sheet(wb)

    wb.save(out)
    return out


# ============================================================================
# Scorecard parser (markdown produced by icos-fit-eval skill)
#
# The parser is permissive — different agents have written gates/LPs in three
# distinct formats. We try multiple patterns per field:
#
#   Inline full   : "Stage/Rev:  UNCLEAR — evidence ..."
#   Inline short  : "Stage/Rev:  U — evidence ..."
#   Table cell    : "| Stage / Revenue | **FAIL** | reason |"
#
# LP scores:
#   Inline bracket: "Nouryon [3/5]"
#   Em-dash       : "Nouryon — 3/5"   or   "### Nouryon — 3/5"
#   Plain         : "Nouryon 3/5"
# ============================================================================
# Each gate has a canonical key + a list of header spellings to try
_GATES = [
    ("g_stage_rev",  ["Stage/Rev",   "Stage / Revenue", "Stage/Revenue", "Stage / Rev"]),
    ("g_sector",     ["Sector fit",  "Sector"]),
    ("g_tech",       ["Technology",  "Tech"]),
    ("g_climate",    ["Climate / CO2", "Climate/CO2", "Climate"]),
    ("g_geography",  ["Geography"]),
    ("g_biz_model",  ["Business model", "Biz model", "Biz Model"]),
    ("g_lp_fit",     ["LP fit (any 3+)", "LP fit", "LP Fit"]),
]
_VALUE_FULL = {"PASS": "P", "FAIL": "F", "UNCLEAR": "U"}


def _gate_value(text: str, names: list[str]) -> str:
    """Try multiple regexes against several name spellings; return P/F/U or empty."""
    for name in names:
        ne = re.escape(name)
        # Inline full word: "Name:  PASS — ..."
        m = re.search(rf"{ne}\s*:\s*\*?\*?(PASS|FAIL|UNCLEAR)\b", text, re.I)
        if m:
            return _VALUE_FULL[m.group(1).upper()]
        # Inline shorthand: "Name:  P — ..."   (single letter followed by — or space)
        m = re.search(rf"{ne}\s*:\s+([PFU])(?:\s+|\s*[—-])", text)
        if m:
            return m.group(1).upper()
        # Table cell: "| Name | **FAIL** |"  or  "| Name | FAIL |"
        m = re.search(rf"\|\s*{ne}\s*\|\s*\*?\*?(PASS|FAIL|UNCLEAR)\*?\*?\s*\|", text, re.I)
        if m:
            return _VALUE_FULL[m.group(1).upper()]
    return ""


def _lp_value(text: str, names: list[str]) -> str:
    """Find the LP's score across the formats agents have produced.

    Tries patterns in order from most-specific to most-permissive."""
    for name in names:
        ne = re.escape(name)
        # 1. "Nouryon — Score: 3/5"
        m = re.search(rf"{ne}[\w\s.,&-]{{0,60}}?[—-]\s*Score\s*:\s*(\d)\s*/\s*5", text, re.I)
        if m:
            return m.group(1) + "/5"
        # 2. bracket form: "Nouryon [3/5]"  (allow qualifier like "Ingredients" between)
        m = re.search(rf"{ne}[\w\s.,&-]{{0,60}}?\[\s*(\d)\s*/\s*5\s*\]", text, re.I)
        if m:
            return m.group(1) + "/5"
        # 3. em-dash form: "Nouryon — 3/5" or "### Nouryon — 3/5"
        m = re.search(rf"{ne}\s*[—-]\s*(\d)\s*/\s*5", text, re.I)
        if m:
            return m.group(1) + "/5"
        # 4. plain space form: "Nouryon 2/5" or "Nouryon  2/5: ..."
        m = re.search(rf"{ne}\s+(\d)\s*/\s*5\b", text, re.I)
        if m:
            return m.group(1) + "/5"
    return ""


def parse_scorecard(text: str) -> dict:
    """Parse an icos-fit-eval scorecard markdown into a flat dict."""
    out: dict = {"score": "", "recommendation": "", "top_question": "", "analyst_note": "",
                 "lp_nouryon": "", "lp_buhler": "", "lp_fc": ""}
    for k, _ in _GATES:
        out[k] = ""

    # Score + recommendation — multiple variants:
    #   "Score: 3/5 | MONITOR"
    #   "**Score: 3/5 | MONITOR**"
    #   "## Score: 1/5 | PASS"
    m = re.search(r"Score\s*:\s*\*?\*?(\d)\s*/\s*5\s*\*?\*?\s*\|\s*\*?\*?(PROCEED|MONITOR|PASS)\*?\*?",
                  text, re.I)
    if m:
        out["score"] = m.group(1) + "/5"
        out["recommendation"] = m.group(2).upper()
    else:
        # Biomatter-style: separate "Composite score: 3/5" + "## Recommendation: MONITOR"
        sm = re.search(r"(?:Composite\s+score|Overall\s+score|Score)\s*:\s*\*?\*?(\d)\s*/\s*5",
                       text, re.I)
        if sm:
            out["score"] = sm.group(1) + "/5"
        rm = re.search(r"#+\s*Recommendation\s*:\s*\*?\*?(PROCEED|MONITOR|PASS)\*?\*?", text, re.I)
        if rm:
            out["recommendation"] = rm.group(1).upper()

    # Gates
    for key, names in _GATES:
        out[key] = _gate_value(text, names)

    # LP scores
    out["lp_nouryon"] = _lp_value(text, ["Nouryon"])
    out["lp_buhler"]  = _lp_value(text, ["Bühler Group", "Bühler", "Buhler"])
    out["lp_fc"]      = _lp_value(text, ["FrieslandCampina Ingredients", "FrieslandCampina",
                                          "Friesland Campina", "FC "])

    # Top question
    qm = re.search(r"##\s*QUESTIONS(.+?)(?:##|$)", text, re.S | re.I)
    if qm:
        first = re.search(r"\d\.\s*\**[^*\n]*?(?:could kill\)?:?\s*\**)?\s*(.+)", qm.group(1))
        if first:
            out["top_question"] = first.group(1).strip().split("\n")[0][:240]

    # Analyst note from RECOMMENDATION (handles `## RECOMMENDATION` and `### RECOMMENDATION`)
    rm = re.search(r"#+\s*RECOMMENDATION\s*(.+?)(?:\n#+\s|\Z)", text, re.S | re.I)
    if rm:
        rec = re.sub(r"\s+", " ", rm.group(1).strip())
        rec = re.sub(r"^\*\*[A-Z]+\*\*\s*[—-]?\s*", "", rec)
        out["analyst_note"] = rec[:500]
    return out


# ============================================================================
# Build final longlist — Phase 5 with Sheet 2 populated
# ============================================================================
def build_final_longlist(slug: str, scorecards: list[dict], *, source_v: int | None = None) -> Path:
    """Take latest longlist-v{n}.xlsx and replace Sheet 2 with structured scorecards.

    scorecards: list of dicts each with at least:
        company, domain, score, recommendation, g_*, lp_*, top_question, analyst_note
        (the keys parse_scorecard returns, plus 'company' and 'domain')
    """
    d = _ensure_run_dir(slug)
    if source_v is None:
        # Find highest-numbered longlist
        candidates = sorted(d.glob("longlist-v*.xlsx"))
        if not candidates:
            raise FileNotFoundError(f"no longlist-v*.xlsx in {d}")
        src = candidates[-1]
    else:
        src = d / f"longlist-v{source_v}.xlsx"
    out = d / f"final-longlist-{slug}.xlsx"

    wb = load_workbook(src)
    if "Icos Fit Scores" in wb.sheetnames:
        del wb["Icos Fit Scores"]
    ws2 = wb.create_sheet("Icos Fit Scores")
    _write_header(ws2, SCORE_HEADERS)

    for i, sc in enumerate(scorecards, start=2):
        row = [
            sc.get("company", ""), sc.get("domain", ""),
            sc.get("score", ""), sc.get("recommendation", ""),
            sc.get("g_stage_rev", "") or sc.get("g_stage", ""),
            sc.get("g_sector", ""), sc.get("g_tech", ""), sc.get("g_climate", ""),
            sc.get("g_geography", "") or sc.get("g_geo", ""),
            sc.get("g_biz_model", "") or sc.get("g_biz", ""),
            sc.get("g_lp_fit", "") or sc.get("g_lp", ""),
            sc.get("lp_nouryon", ""), sc.get("lp_buhler", ""), sc.get("lp_fc", ""),
            sc.get("top_question", ""), sc.get("analyst_note", ""),
        ]
        for j, val in enumerate(row, start=1):
            ws2.cell(row=i, column=j, value=val).alignment = WRAP

        # Color the recommendation cell
        rec = sc.get("recommendation", "")
        c = ws2.cell(row=i, column=4)
        if rec == "PROCEED":   c.fill = GREEN
        elif rec == "MONITOR": c.fill = AMBER
        elif rec == "PASS":    c.fill = RED
        elif rec == "SKIPPED": c.fill = GREY

    widths = [22, 28, 9, 14, 13, 13, 13, 13, 13, 13, 13, 11, 11, 18, 60, 80]
    for j, w in enumerate(widths, start=1):
        ws2.column_dimensions[get_column_letter(j)].width = w
    ws2.row_dimensions[1].height = 30
    for r in range(2, len(scorecards) + 2):
        ws2.row_dimensions[r].height = 90

    wb.move_sheet(ws2, offset=1)  # ensure Longlist sheet is first
    wb.save(out)
    return out


# ============================================================================
# Build final-summary docx (Phase 5)
# ============================================================================
def build_final_summary_docx(slug: str, summary: dict) -> Path:
    """Build runs/[slug]/final-summary-[slug].docx.

    summary dict shape:
      {
        "author": "em@icoscapital.com",
        "date":   "2026-05-09",
        "theme":  "Enzyme design and optimization — refined to ...",
        "stats":  {"total": 97, "evaluated": 11, "skipped": 1,
                   "proceed": 0, "monitor": 4, "pass": 7,
                   "monitor_names": ["..."], "pass_names": ["..."]},
        "top_picks": [
            {"company": "Enzymit", "summary_line": "Israel, Series A $10M, $16.1M total",
             "lp_fit": "FrieslandCampina 4/5 (HMOs/HA), Nouryon 3/5, Bühler 2/5",
             "traction": "World-first multi-kg cell-free HA pilot ...",
             "why_monitor": "Climate gate FAIL (no SFDR LCA) ...",
             "pipeline_note": "Was Lost in Pipedrive Sep 2025 by Kasia ...",
             "retrigger": "EUR 1M+ ARR confirmed + LCA"},
            ...
        ],
        "systemic_findings": [
            "Climate gate hard-failed across the entire cohort. ...",
            "v1/v2 longlist had material quality issues ...",
        ],
        "methodology_lines": ["Round 1 ...", "Round 2 ...", ...],
        "next_step_text": "Reply with names or row numbers; for X and Y loop Kasia in first.",
      }
    """
    from docx import Document
    from docx.shared import Pt

    d = _ensure_run_dir(slug)
    out = d / f"final-summary-{slug}.docx"

    doc = Document()

    # Cover
    doc.add_heading("Herb — Final Report", level=0)
    p = doc.add_paragraph(); p.add_run(f"Run: {slug}").italic = True
    doc.add_paragraph(f"Author: {summary.get('author','')} | Date: {summary.get('date','')}")
    if summary.get("theme"):
        doc.add_paragraph(f"Theme: {summary['theme']}")
    doc.add_paragraph()

    # Summary
    doc.add_heading("Summary", level=1)
    s = summary.get("stats", {})
    doc.add_paragraph(f"• {s.get('total', '?')} companies on the long list (Sheet 1 of the Excel)")
    doc.add_paragraph(f"• {s.get('evaluated', 0)} selected for full Icos Fit evaluation")
    if s.get("skipped"):
        doc.add_paragraph(f"• {s['skipped']} skipped (see scorecard for reason)")
    doc.add_paragraph("• Verdict tally:")
    p = doc.add_paragraph(f"    PROCEED: {s.get('proceed', 0)}"); p.runs[0].bold = True
    p = doc.add_paragraph(f"    MONITOR: {s.get('monitor', 0)}"); p.runs[0].bold = True
    doc.add_paragraph(f"    PASS:    {s.get('pass', 0)}")
    doc.add_paragraph()

    # Systemic findings
    if summary.get("systemic_findings"):
        doc.add_heading("Systemic findings", level=1)
        for i, f in enumerate(summary["systemic_findings"], 1):
            doc.add_heading(f"{i}.", level=2)
            doc.add_paragraph(f)
        doc.add_paragraph()

    # Top picks
    if summary.get("top_picks"):
        doc.add_heading("Top picks (re-engage when triggers met)", level=1)
        for pick in summary["top_picks"]:
            doc.add_heading(pick["company"], level=2)
            p = doc.add_paragraph(); p.add_run(pick.get("summary_line", "")).italic = True
            if pick.get("lp_fit"):
                doc.add_paragraph(f"LP fit: {pick['lp_fit']}")
            if pick.get("traction"):
                doc.add_paragraph(f"Traction: {pick['traction']}")
            if pick.get("why_monitor"):
                doc.add_paragraph(f"Why MONITOR: {pick['why_monitor']}")
            if pick.get("pipeline_note"):
                p = doc.add_paragraph(); p.add_run(f"Pipeline note: {pick['pipeline_note']}").italic = True
            if pick.get("retrigger"):
                p = doc.add_paragraph(); p.add_run(f"Re-engage trigger: {pick['retrigger']}").bold = True
            doc.add_paragraph()

    # Methodology
    if summary.get("methodology_lines"):
        doc.add_heading("Methodology", level=1)
        for line in summary["methodology_lines"]:
            doc.add_paragraph(line)
        doc.add_paragraph()

    # Next step
    if summary.get("next_step_text"):
        doc.add_heading("Next step — Pipedrive entry", level=1)
        doc.add_paragraph(summary["next_step_text"])

    doc.save(out)
    return out


# ============================================================================
# CLI for quick smoke testing
# ============================================================================
if __name__ == "__main__":
    import argparse, json, sys

    ap = argparse.ArgumentParser(description="Herb longlist builder")
    sub = ap.add_subparsers(dest="cmd", required=True)

    p1 = sub.add_parser("v1", help="build longlist-v1.xlsx from a JSON list of rows")
    p1.add_argument("slug"); p1.add_argument("rows_json")

    pn = sub.add_parser("vn", help="build longlist-v{n}.xlsx by appending new_rows JSON")
    pn.add_argument("slug"); pn.add_argument("n", type=int); pn.add_argument("new_rows_json")

    pf = sub.add_parser("final-xlsx", help="build final-longlist-[slug].xlsx from scorecards JSON")
    pf.add_argument("slug"); pf.add_argument("scorecards_json")

    pd = sub.add_parser("final-docx", help="build final-summary-[slug].docx from summary JSON")
    pd.add_argument("slug"); pd.add_argument("summary_json")

    pp = sub.add_parser("parse", help="parse a scorecard markdown file -> JSON")
    pp.add_argument("path")

    args = ap.parse_args()
    if args.cmd == "v1":
        rows = json.loads(Path(args.rows_json).read_text(encoding="utf-8"))
        print(build_longlist_v1(args.slug, rows))
    elif args.cmd == "vn":
        rows = json.loads(Path(args.new_rows_json).read_text(encoding="utf-8"))
        print(build_longlist_vN(args.slug, args.n, rows))
    elif args.cmd == "final-xlsx":
        cards = json.loads(Path(args.scorecards_json).read_text(encoding="utf-8"))
        print(build_final_longlist(args.slug, cards))
    elif args.cmd == "final-docx":
        s = json.loads(Path(args.summary_json).read_text(encoding="utf-8"))
        print(build_final_summary_docx(args.slug, s))
    elif args.cmd == "parse":
        text = Path(args.path).read_text(encoding="utf-8")
        print(json.dumps(parse_scorecard(text), indent=2))
